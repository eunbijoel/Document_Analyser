"""요약/연구노트 → 한글에서 열리는 HWPX·DOCX 생성."""

from __future__ import annotations

import io
from typing import Sequence


def _paragraph_lines(text: str) -> list[str]:
    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    out: list[str] = []
    for ln in lines:
        if ln.strip() == "" and out and out[-1] == "":
            continue
        out.append(ln)
    if not any(x.strip() for x in out):
        return [" "]
    return out


def build_hwpx_from_text(text: str) -> bytes:
    """평문 → HWPX (왼쪽 요약 미리보기용)."""
    from hwpx import HwpxDocument

    doc = HwpxDocument.new()
    lines = _paragraph_lines(text)
    paras = list(getattr(doc, "paragraphs", None) or [])
    if paras:
        try:
            paras[0].text = lines[0]
            start = 1
        except Exception:
            start = 0
    else:
        start = 0
    for ln in lines[start:]:
        doc.add_paragraph(ln if ln.strip() else " ")
    raw = doc.to_bytes()
    if not raw or raw[:2] != b"PK":
        raise RuntimeError("HWPX 생성 실패")
    return raw


def build_docx_from_text(text: str) -> bytes:
    """평문 → DOCX."""
    from docx import Document

    d = Document()
    for ln in _paragraph_lines(text):
        d.add_paragraph(ln)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _normalize_note_rows(rows: Sequence[tuple[str, str]]) -> list[tuple[str, str]]:
    out: list[tuple[str, str]] = []
    for label, val in rows:
        out.append((str(label or ""), str(val or "")))
    if not out:
        out = [("내 용", " ")]
    return out


def _strip_md_noise(text: str) -> str:
    """한글 표 셀용: 마크다운 기호만 가볍게 제거 (내용은 유지)."""
    import re

    out_lines: list[str] = []
    for ln in (text or "").splitlines():
        s = ln.rstrip()
        s = re.sub(r"^#{1,6}\s*", "", s)
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
        s = re.sub(r"__(.+?)__", r"\1", s)
        s = re.sub(r"`([^`]+)`", r"\1", s)
        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", s.strip()):
            continue
        out_lines.append(s)
    return "\n".join(out_lines)


_TALL_LABELS = {"내 용", "연구결과", "기타내용", "내용"}
# A4 한 페이지에 안전하게 들어가는 분량 (셀 단위 나눔은 '행' 단위로만 넘어감)
_MAX_LINES_PER_CELL = 28
_CHARS_PER_LINE = 42


def _visual_line_count(line: str) -> int:
    n = len(line or "")
    if n <= 0:
        return 1
    return max(1, (n + _CHARS_PER_LINE - 1) // _CHARS_PER_LINE)


def _chunk_text_for_pages(text: str, *, max_lines: int = _MAX_LINES_PER_CELL) -> list[str]:
    """한 셀이 페이지보다 커지지 않도록 텍스트를 여러 덩어리로 나눔."""
    raw = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = raw.split("\n") if raw.strip() else [" "]
    chunks: list[str] = []
    buf: list[str] = []
    used = 0
    for ln in lines:
        cost = _visual_line_count(ln)
        if buf and used + cost > max_lines:
            chunks.append("\n".join(buf))
            buf = [ln]
            used = cost
        else:
            buf.append(ln)
            used += cost
    if buf:
        chunks.append("\n".join(buf))
    return chunks or [" "]


def _expand_rows_for_page_breaks(rows: Sequence[tuple[str, str]]) -> list[tuple[str, str]]:
    """긴 내용/결과 행을 여러 표 행으로 쪼개 페이지 넘어가기를 가능하게 함."""
    out: list[tuple[str, str]] = []
    for label, val in rows:
        clean = _strip_md_noise(val)
        key = label.strip()
        compact = key.replace(" ", "")
        if key in _TALL_LABELS or compact in _TALL_LABELS:
            parts = _chunk_text_for_pages(clean)
            for i, part in enumerate(parts):
                out.append((label if i == 0 else "", part))
        else:
            out.append((label, clean if clean.strip() else " "))
    return out


def _configure_hwpx_table_pagination(table) -> None:
    """쪽 경계에서 셀(행) 단위로 넘어가도록 표 속성 보정."""
    el = getattr(table, "element", None)
    if el is None:
        return
    el.set("pageBreak", "CELL")
    el.set("repeatHeader", "0")
    # 글자처럼 취급이면 표가 한 덩어리로 취급되어 페이지 넘김이 깨질 수 있음
    hp = "{http://www.hancom.co.kr/hwpml/2011/paragraph}"
    pos = el.find(f"{hp}pos")
    if pos is not None:
        pos.set("treatAsChar", "0")
        pos.set("flowWithText", "1")
    # 표 전체 고정 높이가 있으면 긴 내용이 잘림 → 최소값만 유지
    sz = el.find(f"{hp}sz")
    if sz is not None:
        try:
            h = int(sz.get("height") or "0")
            if h > 20_000:
                sz.set("height", "20000")
        except Exception:
            pass
    for pos2 in el.iter():
        if str(pos2.tag).endswith("pos") and "treatAsChar" in pos2.attrib:
            pos2.set("treatAsChar", "0")


def _set_row_height(table, row_index: int, height: int) -> None:
    try:
        h = max(2800, min(int(height), 18_000))  # 페이지(~55k)보다 작게 유지
        table.cell(row_index, 0).set_size(height=h)
        table.cell(row_index, 1).set_size(height=h)
    except Exception:
        pass


def build_research_note_hwpx(rows: Sequence[tuple[str, str]], *, title: str = "연구노트") -> bytes:
    """연구노트 2열 표 형식 HWPX (미리보기와 동일 구조)."""
    from hwpx import HwpxDocument

    # python-hwpx 기본 표 폭(14400)은 A4에서 너무 좁음 → authoring 기본(~45000) 사용
    _TABLE_WIDTH = 45_000

    note_rows = _expand_rows_for_page_breaks(_normalize_note_rows(rows))
    doc = HwpxDocument.new()
    paras = list(getattr(doc, "paragraphs", None) or [])
    if paras:
        try:
            paras[0].text = title
        except Exception:
            doc.add_paragraph(title)
    else:
        doc.add_paragraph(title)

    table = doc.add_table(rows=len(note_rows), cols=2, width=_TABLE_WIDTH)
    try:
        # 라벨 ~22% / 값 ~78%
        table.set_column_widths([2200, 7800])
    except Exception:
        pass

    for i, (label, val) in enumerate(note_rows):
        text = val if str(val).strip() else " "
        table.set_cell_text(i, 0, label if label else " ")
        table.set_cell_text(i, 1, text, split_paragraphs=True)
        try:
            table.set_cell_shading(i, 0, "#F0F0F0")
        except Exception:
            pass
        # 행 높이는 내용에 맞추되, 한 페이지를 넘지 않게 상한
        visual = sum(_visual_line_count(ln) for ln in str(text).split("\n"))
        _set_row_height(table, i, 2600 + visual * 750)

    _configure_hwpx_table_pagination(table)

    raw = doc.to_bytes()
    if not raw or raw[:2] != b"PK":
        raise RuntimeError("연구노트 HWPX 생성 실패")
    return raw


def build_research_note_docx(rows: Sequence[tuple[str, str]], *, title: str = "연구노트") -> bytes:
    """연구노트 2열 표 형식 DOCX."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt, RGBColor

    note_rows = _normalize_note_rows(rows)
    d = Document()
    h = d.add_paragraph(title)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h.runs:
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(16)

    table = d.add_table(rows=len(note_rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(3.2)
    table.columns[1].width = Cm(13.0)

    tall_labels = {"내 용", "연구결과", "기타내용", "내용"}

    for i, (label, val) in enumerate(note_rows):
        c0, c1 = table.rows[i].cells
        c0.text = label
        c1.text = _strip_md_noise(val)
        # 라벨 셀 회색 배경
        try:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F0F0F0")
            shading.set(qn("w:val"), "clear")
            c0._tc.get_or_add_tcPr().append(shading)
        except Exception:
            pass
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
        # 내용 행 높이
        if label.strip() in tall_labels or label.replace(" ", "") in {"내용", "연구결과", "기타내용"}:
            try:
                from docx.shared import Twips
                table.rows[i].height = Twips(1800)
            except Exception:
                pass

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _patch_cfb_dir_sector_count_for_v3(hwp_bytes: bytes) -> bytes:
    """hwpkit.cfb.dump가 CFB major v3 파일에 NumberOfDirSectors≠0을 넣는 버그 보정.

    MS-CFB: major version 3이면 offset 0x28(Number of Directory Sectors)는 반드시 0.
    한글/hwpilot은 이 값이 0이 아니면 「파일 손상」으로 거절한다.
    """
    import struct

    if len(hwp_bytes) < 0x30 or hwp_bytes[:4] != b"\xd0\xcf\x11\xe0":
        return hwp_bytes
    data = bytearray(hwp_bytes)
    major = struct.unpack_from("<H", data, 0x1A)[0]
    if major == 3:
        struct.pack_into("<I", data, 0x28, 0)
    return bytes(data)


def build_research_note_hwp(rows: Sequence[tuple[str, str]], *, title: str = "연구노트") -> bytes:
    """한글에서 만든 Note_Template.hwp에 연구노트 필드를 채워 HWP 반환."""
    import shutil
    import tempfile
    from pathlib import Path

    try:
        from hwpkit import open_document
    except ImportError as e:
        raise RuntimeError("HWP 채우기에 hwpkit이 필요합니다. pip install hwpkit") from e

    root = Path(__file__).resolve().parents[1]
    template = root / "templates" / "Note_Template.hwp"
    if not template.is_file():
        raise RuntimeError(f"연구노트 HWP 템플릿 없음: {template}")

    # 라벨 텍스트 → 값 셀(바로 다음 문단). 템플릿 빈 행이 바뀌어도 라벨로 찾음.
    wanted = {
        "주제",
        "책임자",
        "일시",
        "작성자",
        "내용",
        "연구결과",
        "기타내용",
    }
    note_rows = _normalize_note_rows(rows)
    by_label: dict[str, str] = {}
    for label, val in note_rows:
        compact = label.strip().replace(" ", "")
        if compact in wanted:
            clean = _strip_md_noise(val)
            by_label[compact] = clean if clean.strip() else " "

    with tempfile.TemporaryDirectory(prefix="da_hwp_fill_") as tmp:
        work = Path(tmp) / "note.hwp"
        shutil.copy2(template, work)
        doc = open_document(str(work))

        # 템플릿 문단을 훑어 라벨을 찾고, 값은 다음 문단(값 열)에 씀
        values: dict[int, str] = {}
        i = 0
        while True:
            try:
                raw = doc.paragraph_text(i) or ""
            except Exception:
                break
            key = raw.strip().replace(" ", "")
            if key in by_label and (i + 1) not in values:
                values[i + 1] = by_label[key]
            i += 1

        for idx, text in sorted(values.items()):
            try:
                cur = ""
                try:
                    cur = (doc.paragraph_text(idx) or "").strip()
                except Exception:
                    cur = ""
                if cur:
                    doc.replace_text(idx, text)
                else:
                    doc.inject_text(idx, text)
            except Exception:
                try:
                    doc.replace_text(idx, text)
                except Exception:
                    doc.inject_text(idx, text)
        out = Path(tmp) / "research_note.hwp"
        doc.save(str(out))
        data = _patch_cfb_dir_sector_count_for_v3(out.read_bytes())

    if len(data) < 512 or data[:4] != b"\xd0\xcf\x11\xe0":
        raise RuntimeError("생성된 HWP가 손상되었거나 형식이 올바르지 않습니다")
    return data


def build_hwp_from_text(text: str) -> bytes:
    """평문 → HWP (템플릿 '내용' 칸에만 채움)."""
    return build_research_note_hwp([("내 용", text or " ")], title="연구노트")
